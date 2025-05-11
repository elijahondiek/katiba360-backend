#!/usr/bin/env python
# -*- coding: utf-8 -*-
"""
Examine DOCX file to understand its structure
"""

import os
from docx import Document

# Set paths
script_dir = os.path.dirname(os.path.abspath(__file__))
project_dir = os.path.dirname(os.path.dirname(script_dir))
docx_path = os.path.join(project_dir, 'src', 'data', 'source', 'TheConstitutionOfKenya.docx')

# Load document
doc = Document(docx_path)

# Print total paragraphs
print(f'Total paragraphs: {len(doc.paragraphs)}')

# Print first 20 paragraphs
print('\nFirst 20 paragraphs:')
for i, p in enumerate(doc.paragraphs[:20]):
    print(f'{i+1}. [{p.style.name}] {p.text}')

# Print paragraphs with "CHAPTER" in them
print('\nChapter headings:')
for i, p in enumerate(doc.paragraphs):
    if 'CHAPTER' in p.text:
        print(f'{i+1}. [{p.style.name}] {p.text}')

# Print paragraphs with "Article" in them
print('\nArticle headings:')
count = 0
for i, p in enumerate(doc.paragraphs):
    if 'Article' in p.text:
        print(f'{i+1}. [{p.style.name}] {p.text}')
        count += 1
        if count >= 10:  # Limit to first 10 matches
            break
