#!/usr/bin/env python3
"""
DOCX Inspector

A simple script to inspect the content of a DOCX file and print its paragraphs.
"""

from docx import Document
import os
from pathlib import Path
import logging

# Configure logging
logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(name)s - %(levelname)s - %(message)s'
)
logger = logging.getLogger(__name__)

def inspect_docx(docx_path):
    """Inspect a DOCX file and print its contents."""
    try:
        logger.info(f"Opening DOCX file: {docx_path}")
        doc = Document(docx_path)
        
        # Count paragraphs
        total_paragraphs = len(doc.paragraphs)
        non_empty_paragraphs = sum(1 for p in doc.paragraphs if p.text.strip())
        
        logger.info(f"Total paragraphs: {total_paragraphs}")
        logger.info(f"Non-empty paragraphs: {non_empty_paragraphs}")
        
        # Print the first 20 non-empty paragraphs
        logger.info("First 20 non-empty paragraphs:")
        count = 0
        for i, para in enumerate(doc.paragraphs):
            text = para.text.strip()
            if text:
                style_name = para.style.name if para.style else "No style"
                logger.info(f"Paragraph {i} (Style: {style_name}): {text[:100]}")
                count += 1
                if count >= 20:
                    break
        
        # Check for tables
        tables_count = len(doc.tables)
        logger.info(f"Number of tables: {tables_count}")
        
        if tables_count > 0:
            logger.info("First table content:")
            for i, row in enumerate(doc.tables[0].rows[:5]):
                row_text = " | ".join(cell.text for cell in row.cells)
                logger.info(f"Row {i}: {row_text}")
                
    except Exception as e:
        logger.error(f"Error inspecting DOCX file: {e}")
        raise

def main():
    """Main function to inspect the DOCX file."""
    # Get the absolute path to the DOCX file
    base_dir = Path(__file__).resolve().parent.parent.parent
    docx_path = base_dir / "src" / "data" / "source" / "TheConstitutionOfKenya.docx"
    
    # Check if the file exists
    if not docx_path.exists():
        logger.error(f"DOCX file not found at: {docx_path}")
        logger.info(f"Current directory: {os.getcwd()}")
        logger.info(f"Base directory: {base_dir}")
        return
    else:
        logger.info(f"DOCX file found at: {docx_path}")
        logger.info(f"File size: {docx_path.stat().st_size} bytes")
    
    # Inspect the DOCX file
    inspect_docx(str(docx_path))

if __name__ == "__main__":
    main()
